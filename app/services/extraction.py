from pathlib import Path
import csv
import xml.etree.ElementTree as ElementTree
from typing import List, Tuple

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader
from PIL import Image
import pytesseract


# ---------------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------------

def extract_text_from_file(path: str, filename: str | None = None) -> str:
    file_path = Path(path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext in {".txt", ".md", ".log"}:
        return _extract_text(file_path)
    if ext in {".html", ".htm"}:
        return _extract_html(file_path)
    if ext == ".csv":
        return _extract_csv(file_path)
    if ext == ".xml":
        return _extract_xml(file_path)
    if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return _extract_image(file_path)
    if ext in {".xlsx", ".xlsm"}:
        return _extract_excel(file_path)
    if ext == ".xls":
        return _extract_excel_legacy(file_path)

    try:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        name = filename or file_path.name
        raise ValueError(f"Unsupported file type for {name}") from exc


# ---------------------------------------------------------------------------
# Existing extractors
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs if p.text)
        parts.extend(p.text for p in section.footer.paragraphs if p.text)

    return "\n".join(parts)


def _extract_text(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


def _extract_html(file_path: Path) -> str:
    raw = file_path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")
    return soup.get_text(separator=" ", strip=True)


def _extract_csv(file_path: Path) -> str:
    lines = []
    with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for idx, row in enumerate(reader):
            if idx >= 150:
                break
            lines.append("\t".join(row))
    return "\n".join(lines)


def _extract_xml(file_path: Path) -> str:
    tree = ElementTree.parse(str(file_path))
    root = tree.getroot()
    texts = []
    for element in root.iter():
        if element.text and element.text.strip():
            texts.append(element.text.strip())
    return "\n".join(texts)


def _extract_image(file_path: Path) -> str:
    try:
        image = Image.open(str(file_path))
        return pytesseract.image_to_string(image)
    except Exception as exc:
        raise RuntimeError(f"OCR failed: {exc}")


# ---------------------------------------------------------------------------
# Excel helpers
# ---------------------------------------------------------------------------

def _detect_table_blocks(
    rows: List[List[str]],
) -> List[Tuple[int, int]]:
    """Return (start_row, end_row) index pairs for each contiguous data block.

    A *block* is a maximal sequence of non-blank rows.  Two rows with at least
    one non-empty cell are separated into different blocks by one or more rows
    that are entirely blank.

    Row indices are 0-based and *inclusive* on both ends.
    """
    blocks: List[Tuple[int, int]] = []
    start: int | None = None

    for idx, row in enumerate(rows):
        is_blank = all(str(cell).strip() == "" for cell in row)
        if not is_blank:
            if start is None:
                start = idx
        else:
            if start is not None:
                blocks.append((start, idx - 1))
                start = None

    if start is not None:
        blocks.append((start, len(rows) - 1))

    return blocks


def _rows_to_text(rows: List[List[str]]) -> str:
    """Convert a 2-D list of strings into tab-separated lines."""
    return "\n".join("\t".join(cell for cell in row) for row in rows)


def _extract_excel(file_path: Path) -> str:
    """Extract text from .xlsx / .xlsm using openpyxl (no xlrd dependency)."""
    import openpyxl  # lazy import — only needed for Excel files

    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    sections: List[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # Collect all rows as lists of strings, preserving empty cells
        raw_rows: List[List[str]] = []
        for row in ws.iter_rows(max_row=150, values_only=True):
            raw_rows.append([str(cell) if cell is not None else "" for cell in row])

        # Trim trailing all-blank columns across all rows
        if raw_rows:
            max_col = max(
                (
                    max(
                        (i + 1 for i, c in enumerate(row) if c.strip()),
                        default=0,
                    )
                    for row in raw_rows
                ),
                default=0,
            )
            raw_rows = [row[:max_col] for row in raw_rows]

        blocks = _detect_table_blocks(raw_rows)
        for block_idx, (start, end) in enumerate(blocks, start=1):
            block_rows = raw_rows[start : end + 1]
            # Human-readable 1-based row numbers matching the spreadsheet
            label = (
                f"[Sheet: {sheet_name} | Table {block_idx} "
                f"(rows {start + 1}–{end + 1})]"
            )
            sections.append(f"{label}\n{_rows_to_text(block_rows)}")

    return "\n\n".join(sections)


def _extract_excel_legacy(file_path: Path) -> str:
    """Extract text from .xls using pandas + xlrd."""
    try:
        import pandas as pd  # lazy import
    except ImportError as exc:
        raise RuntimeError(
            "pandas is required to read .xls files. "
            "Install it with: pip install pandas xlrd"
        ) from exc

    sections: List[str] = []
    xls = pd.ExcelFile(str(file_path), engine="xlrd")

    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, header=None, nrows=150, dtype=str).fillna("")
        raw_rows: List[List[str]] = df.values.tolist()  # type: ignore[assignment]

        blocks = _detect_table_blocks(raw_rows)
        for block_idx, (start, end) in enumerate(blocks, start=1):
            block_rows = raw_rows[start : end + 1]
            label = (
                f"[Sheet: {sheet_name} | Table {block_idx} "
                f"(rows {start + 1}–{end + 1})]"
            )
            sections.append(f"{label}\n{_rows_to_text(block_rows)}")

    return "\n\n".join(sections)
